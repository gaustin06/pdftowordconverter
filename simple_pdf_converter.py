#!/usr/bin/env python3

import os
import sys
import logging
from pathlib import Path
from typing import List, Optional

try:
    import pdfplumber
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError as e:
    print("=" * 60)
    print("❌ ERROR: Missing dependencies")
    print("=" * 60)
    print("Run the following command in your terminal:")
    print("pip install pdfplumber python-docx")
    print("=" * 60)
    sys.exit(1)


class SimplePDFConverter:
    """Simplified PDF to Word converter for Seismic LiveDocs"""
    
    def __init__(self, output_folder: str, verbose: bool = True):
        self.output_folder = Path(output_folder)
        self.verbose = verbose
        self.setup_logging()
        self.create_output_folder()
        
    def setup_logging(self):
        """Configure logging system"""
        level = logging.INFO if self.verbose else logging.ERROR
        
        logging.basicConfig(
            level=level,
            format='%(asctime)s - %(message)s',
            handlers=[
                logging.FileHandler('conversion.log'),
                logging.StreamHandler(sys.stdout)
            ]
        )
        self.logger = logging.getLogger(__name__)
        
    def create_output_folder(self):
        """Create output folder"""
        try:
            self.output_folder.mkdir(parents=True, exist_ok=True)
            self.logger.info(f"📁 Output folder: {self.output_folder}")
        except Exception as e:
            print(f"❌ Error creating folder: {e}")
            raise
            
    def validate_pdf_files(self, pdf_files: List[str]) -> List[str]:
        """Validate and filter existing PDF files"""
        valid_pdfs = []
        
        print("🔍 Validating PDF files...")
        
        for pdf_file in pdf_files:
            pdf_path = Path(pdf_file)
            
            if not str(pdf_file).lower().endswith('.pdf'):
                print(f"⚠️  Ignored (not PDF): {pdf_file}")
                continue
                
            if not pdf_path.exists():
                print(f"❌ Not found: {pdf_file}")
                continue
                
            valid_pdfs.append(pdf_file)
            print(f"✅ Valid: {pdf_file}")
            
        return valid_pdfs
        
    def extract_text_from_pdf(self, pdf_path: str) -> List[dict]:
        """Extract text from PDF"""
        pages_content = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    
                    if text:
                        pages_content.append({
                            'page_number': i + 1,
                            'text': text.strip(),
                            'tables': page.extract_tables() or []
                        })
                        
            self.logger.info(f"📄 Extracted {len(pages_content)} pages from {Path(pdf_path).name}")
            return pages_content
            
        except Exception as e:
            self.logger.error(f"❌ Error extracting text: {e}")
            return []
    
    def has_meaningful_content(self, table_data: List[List[str]]) -> bool:
        """Check if table has meaningful content"""
        if not table_data:
            return False
            
        # Count cells with content
        cells_with_content = 0
        total_cells = 0
        
        for row in table_data:
            for cell in row:
                total_cells += 1
                if cell and str(cell).strip():
                    cells_with_content += 1
        
        # Only consider tables with at least 30% content
        return total_cells > 0 and (cells_with_content / total_cells) >= 0.3
    
    def process_text_content(self, doc: Document, text_lines: List[str]):
        """Process text content intelligently"""
        current_paragraph = None
        skip_empty_lines = 0
        
        for i, line in enumerate(text_lines):
            line = line.strip()
            
            # Skip completely empty lines
            if not line:
                if current_paragraph is not None:
                    current_paragraph = None
                continue
            
            # Skip lines that appear to be page numbers alone
            if line.isdigit() and len(line) <= 3:
                continue
                
            # Detect main titles
            if self.is_main_title(line):
                doc.add_heading(line, level=1)
                current_paragraph = None
                continue
            
            # Detect subtitles
            if self.is_subtitle(line):
                doc.add_heading(line, level=2)
                current_paragraph = None
                continue
                
            # Detect list elements
            if self.is_bullet_point(line):
                # Clean bullet point for cleaner format
                clean_line = self.clean_bullet_point(line)
                doc.add_paragraph(clean_line, style='List Bullet')
                current_paragraph = None
                continue
            
            # Normal text
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            
            # Add text
            if current_paragraph.text:  # If paragraph already has text
                current_paragraph.add_run(' ' + line)
            else:
                current_paragraph.add_run(line)
            
            # End paragraph in certain cases
            if line.endswith('.') or line.endswith(':') or line.endswith('!') or line.endswith('?'):
                # Check if next line suggests a new paragraph
                next_line = text_lines[i + 1].strip() if i + 1 < len(text_lines) else ""
                if (not next_line or 
                    self.is_main_title(next_line) or 
                    self.is_subtitle(next_line) or 
                    self.is_bullet_point(next_line) or
                    next_line[0].isupper() if next_line else False):
                    current_paragraph = None
    
    def is_main_title(self, line: str) -> bool:
        """Detect main titles"""
        if not line or len(line) > 150:
            return False
            
        # Short titles in uppercase
        if len(line) < 80 and line.isupper():
            return True
            
        # Lines ending without punctuation and relatively short
        if (len(line) < 100 and 
            not line.endswith('.') and 
            not line.endswith(',') and
            line[0].isupper() and
            any(word in line.lower() for word in [
                'overview', 'introduction', 'benefits', 'features', 'solution', 
                'services', 'why', 'how', 'what', 'branded calling', 'productivity',
                'appendix', 'methodology', 'verizon', 'empowerment'
            ])):
            return True
            
        return False
    
    def is_subtitle(self, line: str) -> bool:
        """Detect subtitles"""
        if not line or len(line) > 120:
            return False
            
        # Lines starting with capital and are question or descriptive
        if (len(line) < 120 and 
            line[0].isupper() and
            (line.endswith('?') or 
             (not line.endswith('.') and not line.endswith(',')) and
             any(word in line.lower() for word in [
                 'create', 'project', 'customer', 'workshop', 'deliverables',
                 'findings', 'different', 'calling', 'profiles', 'answer', 'quandary'
             ]))):
            return True
            
        return False
    
    def clean_bullet_point(self, line: str) -> str:
        """Clean bullet points for consistent format"""
        # Remove common markers
        for marker in ['•', '-', '*', '○']:
            if line.startswith(marker):
                return line[1:].strip()
        
        # Remove numbering
        import re
        if re.match(r'^\d+\.', line):
            return re.sub(r'^\d+\.\s*', '', line)
            
        return line
            
    def create_word_document(self, content: List[dict], title: str) -> Document:
        """Create clean Word document following example format"""
        doc = Document()
        
        # Configure styles
        self.setup_document_styles(doc)
        
        # Process all content as continuous document
        all_text_lines = []
        all_tables = []
        
        # Combine all text from all pages
        for page_info in content:
            page_lines = page_info['text'].split('\n')
            all_text_lines.extend(page_lines)
            
            # Collect tables
            for table_data in page_info.get('tables', []):
                if table_data and self.has_meaningful_content(table_data):
                    all_tables.append(table_data)
        
        # Process text lines intelligently
        self.process_text_content(doc, all_text_lines)
        
        # Add tables at the end if any
        for table_data in all_tables:
            self.add_table_to_document(doc, table_data)
                    
        return doc
        
    def setup_document_styles(self, doc: Document):
        """Configure styles for Seismic LiveDocs"""
        try:
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            
            # Custom style for LiveDocs
            try:
                livedoc_style = doc.styles.add_style('LiveDoc Variable', WD_STYLE_TYPE.CHARACTER)
                livedoc_style.font.name = 'Calibri'
                livedoc_style.font.size = Pt(11)
                livedoc_style.font.bold = True
            except:
                pass
                
        except Exception as e:
            self.logger.warning(f"⚠️ Error configuring styles: {e}")
        
    def is_bullet_point(self, line: str) -> bool:
        """Detect bullet points"""
        return line.startswith(('•', '-', '*', '○')) or line.startswith(tuple(f'{i}.' for i in range(10)))
        
    def add_table_to_document(self, doc: Document, table_data: List[List[str]]):
        """Add table to document with better formatting"""
        if not table_data or not table_data[0]:
            return
            
        try:
            # Clean table data
            clean_data = []
            for row in table_data:
                clean_row = []
                for cell in row:
                    if cell is not None:
                        clean_cell = str(cell).strip()
                        clean_row.append(clean_cell)
                    else:
                        clean_row.append('')
                clean_data.append(clean_row)
            
            # Check if we have useful content
            if not any(any(cell for cell in row) for row in clean_data):
                return
            
            # Create table
            max_cols = max(len(row) for row in clean_data) if clean_data else 0
            if max_cols == 0:
                return
                
            table = doc.add_table(rows=len(clean_data), cols=max_cols)
            table.style = 'Table Grid'
            
            # Fill data
            for i, row_data in enumerate(clean_data):
                for j, cell_data in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = cell_data
            
            # Add space after table
            doc.add_paragraph()
                        
        except Exception as e:
            self.logger.warning(f"⚠️ Error adding table: {e}")
            
    def detect_image_placeholders(self, text_lines: List[str]) -> List[int]:
        """Detect possible image locations based on spaces or references"""
        image_positions = []
        
        for i, line in enumerate(text_lines):
            line_lower = line.lower().strip()
            
            # Look for image or figure references
            if any(keyword in line_lower for keyword in [
                'figure', 'image', 'logo', 'chart', 'graph', 'diagram',
                'see below', 'shown above', 'illustration'
            ]):
                image_positions.append(i)
                
            # Detect very short lines that could be image titles
            if (len(line.strip()) < 50 and 
                any(word in line_lower for word in ['fig', 'image', 'logo'])):
                image_positions.append(i)
        
        return image_positions
        
    def add_image_placeholder(self, doc: Document, placeholder_text: str = "[IMAGE]"):
        """Add image placeholder"""
        p = doc.add_paragraph()
        run = p.add_run(placeholder_text)
        run.font.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    def convert_pdf_to_word(self, pdf_path: str) -> Optional[str]:
        """Convert a PDF to Word"""
        try:
            content = self.extract_text_from_pdf(pdf_path)
            if not content:
                print(f"❌ Could not extract content from {pdf_path}")
                return None
                
            # Generate Word filename
            pdf_name = Path(pdf_path).stem
            word_filename = f"{pdf_name}.docx"
            word_path = self.output_folder / word_filename
            
            # Create document - no automatic title
            doc = self.create_word_document(content, pdf_name)
            
            # Minimal metadata
            doc.core_properties.comments = "Converted for Seismic LiveDocs use"
            
            # Save
            doc.save(str(word_path))
            print(f"✅ Converted: {Path(pdf_path).name} → {word_filename}")
            
            return str(word_path)
            
        except Exception as e:
            print(f"❌ Error converting {pdf_path}: {e}")
            return None
            
    def convert_all(self, pdf_files: List[str]) -> dict:
        """Convert all PDFs"""
        # Validate files
        valid_pdfs = self.validate_pdf_files(pdf_files)
        
        if not valid_pdfs:
            print("❌ No valid PDF files found")
            return {'successful': 0, 'failed': 0, 'files': []}
            
        print(f"\n🚀 Starting conversion of {len(valid_pdfs)} files...")
        print("-" * 50)
        
        results = {'successful': 0, 'failed': 0, 'files': []}
        
        for pdf_file in valid_pdfs:
            word_file = self.convert_pdf_to_word(pdf_file)
            
            if word_file:
                results['successful'] += 1
                results['files'].append({
                    'pdf': pdf_file,
                    'word': word_file,
                    'status': 'success'
                })
            else:
                results['failed'] += 1
                results['files'].append({
                    'pdf': pdf_file,
                    'word': None,
                    'status': 'failed'
                })
                
        return results